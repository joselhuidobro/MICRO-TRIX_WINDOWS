from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
import pandas as pd
from werkzeug.datastructures import ImmutableMultiDict
from datetime import datetime
from flask import current_app
from app.utils.prediction import predict_location
from app.utils.correo2 import correo2
import logging
import uuid
import logging


logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler()  
    ]
)

logger = logging.getLogger(__name__)
def cotizar_trix(form_data: ImmutableMultiDict):


    
    # Configurar logging
    logging.basicConfig(level=logging.INFO)
    logger.info("Data recibida por el scropt cotizar_trix: %s", form_data)


    document = Document()
    header = document.sections[0].header
    header.is_linked_to_previous = False
    header.center_header = True

    # Obtener la ruta de la carpeta "generados"
    image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'encabezado.jpg')

    logo_paragraph = header.add_paragraph()
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo = logo_paragraph.add_run()
    logo.add_picture(image_path, width=Inches(2.5))



    # Convertir el ImmutableMultiDict en una lista de tuplas, excluyendo el elemento 'csrf_token', quitar el token de la tupla de tuplas
    form_data_list = [(k, v) for k, v in form_data.items() if k != 'csrf_token']

    # Crear el DataFrame a partir de form_data_list
    df_cotizacion = pd.DataFrame(form_data_list, columns=['Campo', 'Valor'])
    # Extraer los productos del formulario
    productos_form = [v for k, v in form_data_list if 'producto' in k.lower()]
    Nombre_cliente = df_cotizacion.loc[df_cotizacion['Campo'] == 'nombre', 'Valor'].iloc[0]
    Direccion_cliente = df_cotizacion.loc[df_cotizacion['Campo'] == 'direccion', 'Valor'].iloc[0]
    Ciudad_cliente = "Ciudad de México" 
    Ciudad_TRIX = "Ciudad de México"  

    fecha_cotizacion = datetime.today().strftime('%Y-%m-%d')

    intro_text = 'Por la presente le mandamos un cordial saludo y tenemos el gusto de someter a su atenta consideración el siguiente presupuesto.'
    body_text = 'Suministro e instalación de sistema abre-puerta con equipo a base de brazos electromecánicos modelo Power 230. De alimentación 24 vcd. Con capacidad de 200 kg por hoja y una frecuencia de uso de 30 activaciones por hora. Cuenta con llave de desbloqueo, 4 telemandos digitales para automóvil (dos por par de brazos), tope para impacto de puerta y lo necesario para su funcionamiento, batería de respaldo. Instalado en puerta de dos hojas abatibles.'
    body_text_roller = 'Suministro e instalación de abre-puertas semi-industrial marca Merik modelo Roller de alimentación 110 Vac, cuenta con motor a 24Vdc, para el acceso de vehículos a través de una cancela batiente. La función de anti-aplastamiento está garantizada por un dispositivo electrónico instalado en la tableta de control que regula directamente el paro del motor. Se instala soterrado en el terreno, por lo que no altera la estética de la cancela. Cuenta con brazo de desbloqueo para liberar la puerta del sistema automático en caso de ausencia de suministro eléctrico. Incluye caja de cimentación, Tina, dos operadores, llave de desbloqueo, un control a distancia, centro de control electrónico. Largo máximo de hoja: 2 – 2.5 metros. Peso por hoja: 500 kg.' 
    body_text_power230 = 'Suministro e instalación de sistema abre-puerta con equipo a base de un par de brazos electro-mecánicos modelo Power 230. De alimentación a  110 vac y motor @ 24 vcd. Con capacidad de 200 kg por hoja y una frecuencia de uso de 30 activaciones por hora. Cuenta con llave de desbloqueo, 2 telemandos digitales para automóvil, batería de respaldo y lo necesario para su correcto funcionamiento, cable de uso rudo 18 AWG para alimentación de actuadores. Para automatizar acceso de dos hojas abatibles al exterior.'
    body_text_gbat110 = 'Suministro e instalación de un sistema de apertura de puertas, equipado con brazos electromecánicos de la marca Genius (Italiana), modelo Gbat a 110 V CA. El sistema incluye los siguientes componentes: un par de brazos electromecánicos, un centro de control (cerebro), un receptor de control remoto, dos controles remotos digitales para automóviles, un botón de apertura, una llave de desbloqueo, herrajes y todos los elementos necesarios para su correcto funcionamiento.'
    body_text_gbat24v = 'Suministro e instalación de un sistema de apertura de puertas, equipado con brazos electromecánicos de la marca Genius, modelo Gbat a 24dc. El sistema incluye los siguientes componentes: un par de brazos electromecánicos, un centro de control (cerebro), un receptor de control remoto, dos controles remotos digitales para automóviles, un botón de apertura, una llave de desbloqueo, Baterías de respaldo, herrajes y todos los elementos necesarios para su correcto funcionamiento.'
    
    body_text_711 = 'Tenemos el gusto de someter a su atenta consideración el siguiente presupuesto. Suministro e instalación de abre puertas residencial a base de cadena marca Merik-Pro modelo 711-M de 1/2 h.p. a 24 vcd. Alimentación a 115 vca. , monofásico con motor reductor de engranes y transmisión de cadena link-belt con corredera en el sentido de apertura de la puerta y brazos articulados de acople a la misma, además cuenta con una luz de encendido de 4 1/2  min. De apagado automático, dos telemandos digitales para automóvil y un botón para abrir desde el interior .' 
    body_text_511 = 'Suministro e instalación de  abre puertas residencial a base de cadena marca Merik-Pro modelo 511-M de 1/2 h.p.  Alimentación a 115 vca. , monofásico con motorreductor de engranes y transmisión de cadena link-belt con corredera en el sentido de apertura de la puerta y brazos articulados de acople a la misma, ademas cuenta con una luz de encendido de 4 1/2  min. de apagado automático ,dos telemandos digitales para automóvil c/u. y un botón para abrir desde el interior .' 
    body_text_fotocelda = 'Suministro e instalación de fotoceldas para puerta automatica, este sensor permite evitar colisiones con personas y objetos que puedan encontrarse en el camino de la puerta, aumentando significativamente la seguridad en tu hogar o negocio.' 
    body_text_pivus30024v ='Suministro e instalación de sistema abre-puerta con equipo a base de par de brazos electro-mecánicos marca SEG modelo Pivus 300 . De alimentación a 24 vcd. Con capacidad de 200 kg por hoja y una frecuencia de uso de 30 activaciones por hora. Cuenta con llave de desbloqueo, 2 telemandos digitales para automóvil, batería de respaldo y lo necesario para su correcto funcionamiento.'
    body_text_cadenaCh = 'Ajuste de cadena chica para puerta automática con mecanismo de Cadena'
    body_text_cadenaG = 'Ajuste de cadena Grande para puerta automática con mecanismo de Cadena'
    body_text_carretilla = 'Suministro e instalacion de carretilla para puerta de acceso vehicular'
    body_text_mantenimiento = 'Servicio de mantenimiento para puerta automática. Incluye limpieza y lubricación de rieles, ajuste mecánico y parámetros electrónicos para el correcto funcionamiento del sistema. Se recomienda dar servicio de mantenimiento cada 6 meses para disminuir el desgaste de los elementos mecánicos. '
    body_text_mantenimiento_hidraulico = 'Servicio de mantenimiento para acceso vehicular con puertas automáticas equipadas con brazos hidráulicos. Esta cotización incluye el cambio de fluido no compresible para el pistón, sustitución de empaques, ajuste mecánico y configuración de parámetros electrónicos para optimizar la apertura y cierre del sistema. Cabe mencionar que no se incluye el cambio de componentes electrónicos, elementos mecánicos ni el reemplazo del ensamble de válvulas.'
    body_text_mantenimiento_cerca = 'Servicio de mantenimiento integral para Cerca Electrificada de Seguridad, incluyendo el ajuste de tensión de cable electrificado, aseguramiento y corrección de la alineación de postes, inspección y calibración de la tarjeta electrónica de control, además de verificar el funcionamiento óptimo de la sirena y otros dispositivos periféricos.'
    body_text_trabajo_herreria = 'Tipo de puerta: [tipo de puerta a fabricar, por ejemplo: corredera, abatible, plegable, etc.]  Herrajes: [detalle de los herrajes necesarios para la puerta, por ejemplo: cerraduras, manijas, bisagras, etc.] Instalación: [detalle de la instalación de la puerta, incluyendo si se requiere la colocación de marcos, ajustes y nivelación]'
    body_text_bft110v='Suministro e instalación de champa electromecánica BFT @110vac con activación automática de tableta de control de pistones'
    body_text_chapa24v='Suministro e instalación de champa electromecánica  @24vdc con activación automática con operacion de motores'
    body_text_cremallera='Suministro e instalación de cremallera de acero inoxidable para puerta corrediza con Motor de Piñón-Cremallera'    
    body_text_tarjeta_mec200='Ofrecemos el suministro, instalación y programación de la tarjeta de control electrónica modelo Mec 200 para el actuador de puerta automática marca Fadini. Esta tarjeta es compatible con el modelo Mec 200.'
    body_text_5011='Suministro e instalación del operador semi industrial de puertas modelo Merik 5011. Este sistema se distingue por su motor de 1/2 h.p. a 120 VCA, monofásico, con un sistema de cadena. Posee un sistema de códigos rodantes para una mayor seguridad. Además, incluye una luz de cortesía que se enciende automáticamente cuando se opera el motor. El paquete comprende un control remoto para el automóvil y un botón de pared para abrir la puerta desde el interior de su hogar. Este motor cuenta con un mecanismo de embrague diseñado para prevenir aplastamientos.'
    body_text_shoker='Suministro e instalación de una cerca electrificada de la marca SEG, modelo SHOCKER 14000 W, diseñada para cubrir un perímetro determinado. El sistema viene equipado con un gabinete que alberga un circuito de control, el cual opera con una alimentación de 110 Vac para la cerca eléctrica y una alarma de emergencia. Además, cuenta con un control de frecuencia de 433 Mhz.'
    body_text_Dahua_2mp='Ofrecemos el suministro e instalación de cuatro cámaras Dahua de tipo domo, con una resolución de 2 megapíxeles y lente de 2.8mm. Cada cámara tiene un campo de visión de 100 grados para una amplia cobertura. El servicio incluye todo el cableado necesario hasta la ubicación del DVR y otros componentes electrónicos. Nuestro equipo profesional garantizará una instalación eficiente y una configuración óptima del sistema.'
    body_text_Dahua_PZT='Ofrecemos el suministro e instalación de la cámara PTZ antivandálica DAHUA SD50225-HC-LA. Esta avanzada cámara de seguridad cuenta con una resolución de 2 Megapixeles y ofrece un zoom óptico de 25x. Su tecnología Starlight y WDR real de 120 dB garantizan una visibilidad clara y nítida incluso en condiciones de iluminación difíciles. El equipo cumple con las normativas IP66 e IK10, lo que confirma su robustez y capacidad para resistir condiciones adversas. Con funciones BLC y HLC y un giro continuo de 360 grados, la cámara ofrece una cobertura completa y confiable de la zona de vigilancia.'
    body_text_Seg_2puertas_controlacceso='Suministro e instalación de una tarjeta de control de accesos Marca SEG que funciona a través de tags RFID y antenas para la identificación de vehículos. El paquete incluye dos antenas con un alcance de 15 metros, 100 tarjetas marca SEG para gestionar hasta 2 puertas y que cuenta con comunicación IP y Wiegand. Esta tarjeta  estará conectada al actuador de la puerta , que se cotizaron previamente. Como parte de este paquete, también se incluyen 100 tags RFID y la instalación del software correspondiente desde la central. Esto permite una gestión eficaz de las tarjetas y accesos si existe un operador dedicado para esta tarea. Las características técnicas de la tarjeta de control de acceso son las siguientes: '
    body_text_Faac_400_CBACL_kit='Se realizará la instalación de un par de Brazos Hidráulicos FAAC modelo 400 de 110vAc de alimentación en un accesos vehicular. El acceso vehicular contará con dos actuadores, una tarjeta de control y una chapa electromecánica, siendo esta última operada desde el circuito de control. Los actuadores están diseñados para puertas de hasta 4m con chapa y 2.5m sin ella. Incluye un botón de apertura y cierre ( uno por acceso) instalado al interior de la portería a 40m de distancia de la puerta de acceso.    '
    body_text_Faac_350h = 'Se realizará la instalación de un par de Brazos Hidráulicos FAAC modelo 350 de 110vAc de alimentación en un accesos vehicular. El acceso vehicular contará con dos actuadores, una tarjeta de control . Los actuadores están diseñados para puertas de hasta 4m con chapa y 3.5 sin ella. Incluye un botón de apertura y cierre ( uno por acceso) instalado al interior de la portería .'
    body_text_FAAC_455D=''
    body_text_tejuelosybibel='Este servicio incluye el suministro e instalación de tejuelos y bibeles, que se instalarán en los extremos superiores e inferiores de la apertura total de acceso. Para poder realizar estos cambios, será necesario desmontar las puertas. De esta manera, las puertas podrán abrir y cerrar con mayor facilidad.'
    body_text_pila12v='Suminitro e instalacion de dos baterias de resplado marca SEG para motor de puerta automática '
    body_text_SEGDuoCombat = " servicio de suministro e instalación del par de brazos marca Seg modelo Duo Combat de 110 VAC de alimentación. Este paquete incluye dos controles remotos, herrajes y todos los elementos necesarios para garantizar un correcto funcionamiento del sistema."
    body_text_dashcam = '''Suministro e instalación de DashCam para vehículos. 
                    🎥 Ofrece captura de videos de alta resolución en 1600p y un amplio campo de visión de 130°.
                    🔌 Es Plug and Play: Se enciende y comienza a grabar automáticamente al arrancar tu coche.
                    💾 Cuenta con capacidad de memoria para tarjetas micro SD de hasta 128 GB para guardar tus videos (tarjeta no incluida).'''

    boy_text_doormakaba80kg  = """ Suministro e instalación de cierra puertas marca Dormakaba para puertas peatonales con un peso máximo de 80 kg y dimensiones de ancho entre 0.85 m y 1.1 m."""
    body_text_merik411= "Suministro e instalación de un abre puertas residencial de cadena de la marca Merik, modelo 411-M de 1/2 h.p. Alimentación a 115 VCA, monofásico, con motorreductor de engranajes y transmisión de cadena link-belt con corredera en el sentido de apertura de la puerta y brazos articulados para acoplarse a la misma. Además, cuenta con una luz que se enciende por 4 1/2 minutos y se apaga automáticamente, módulo Wifi, dos mandos a distancia digitales para automóviles y un botón para abrir desde el interior."                 
    body_text_merikMYQ= "Suministro e instalación de un abre puertas residencial de cadena de la marca Liftmaster-Merik, modelo MYQ de 1/2 h.p. Alimentación a 115 VCA, motor a  monofásico, con motorreductor de engranajes y transmisión de cadena link-belt con corredera en el sentido de apertura de la puerta y brazos articulados para acoplarse a la misma. Además, cuenta con una luz que se enciende por 4 1/2 minutos y se apaga automáticamente, módulo Wifi, dos mandos a distancia digitales para automóviles y un botón para abrir desde el interior."                 
  
    body_text_fococelda_power230 = "Se proporcionará e instalará una fotocelda de seguridad para la puerta automática, esencial para activar el cierre automático tras un periodo determinado. Esta fotocelda es compatible con la tarjeta del motor Merik Power230. Posee un voltaje de 12-24V DC/24 V AC, un rango de alcance de 7 metros y un LED indicador de alineación y suciedad."
    body_text_videoporterohikvision ="El biométrico MarkaHikvision modelo Terminal Min Moe con reconocimiento facial y video portero ofrece una solución avanzada de identificación. Este sistema cuenta con múltiples características, incluyendo la lectura de huellas dactilares para brindar una identificación precisa y segura. Además, permite el acceso mediante tarjetas de identificación RFID, ofreciendo una alternativa cómoda para los usuarios. La comunicación con un software de gestión de acceso facilita la administración centralizada y la reprogramación de visitantes de manera eficiente. Con capacidad para almacenar y reconocer hasta 1500 rostros y 3000 tarjetas de identificación, es ideal para entornos con una alta cantidad de usuarios. Esta solución integral garantiza un control de acceso confiable y seguro para diversos entornos."
    body_text_receptor433hz = "Suministro e instalación de receptor de 433Hz para controles de radiofrecuencia de puerta automática."
    body_text_receptoradministrableSEG = "Presentamos el Receptor Administrable de la marca SEG, una solución de alta calidad y fácil de usar que opera en la frecuencia de 433,92 MHz. Este dispositivo le permite gestionar fácilmente los controles remotos, permitiéndole agregar o eliminar controles de la memoria del receptor con total comodidad."
    body_text_contol_llavero_433 = "Suministro y programación de control tipo llavero de cuatro botones de 433Hz para puerta automática."
    body_text_sirena_12v_110db_600ma_sfire = "Suministro e instalación de sirena marca Sfire, 12VDC, 600mA, 110dB."
    body_text_kit_sirena_control= 'Suministro e instalación de receptor de 433Hz para controles de radiofrecuencia de puerta automática, así como el suministro y programación de N controles tipo llavero de cuatro botones de 433Hz para activar sirena Marca Sfire de 12v de alimentación, 600ma y 110db.'
    body_text_DVR16chan = "Ofrecemos el suministro e instalación de un DVR de 16 canales de la marca Dahua, capaz de conectar hasta 16 cámaras de seguridad. Utilizaremos 50 metros de cable UTP para una transmisión de señal confiable y sin interferencias. Además, incluimos una fuente de poder con multicontacto para la alimentación de las cámaras, evitando así problemas de señal por caídas de tensión. Proporcionamos también los conectores de alimentación y baluns necesarios para la conexión de las cámaras, los cuales convierten señales no balanceadas en señales balanceadas para una transmisión óptima a través de cables UTP. "
    body_text_afpamsuper1000="Suministro e instalación de Motor abre puertas residencial, para actuar domo proporcionado por el cliente. Con mecanismo a base de cadena y Motor marca Afpam modelo SUPER 1000 de Alimentación a 115 VCa., motor @24V dc, con motor reductor de engranes y transmisión de cadena con corredera en el sentido de apertura del Domo, cuenta con una luz de encendido. De apagado automático, 2 controles remotos digitales para automóvil y botón para abrir desde el interior.  "
    body_text_Tensor_motorcadena = "Ofrecemos suministro e instalación de tensor con cuerda para ajustar el cierre de puertas con motor de cadena. Este elemento mejora el cierre al adaptar la distancia de los brazos acoplados a la puerta."
    body_text_duelas_madera= "El proyecto implica suministrar e instalar duelas de madera para cubrir un área total de ______m² en la puerta de acceso vehicular. Esta área no incluye la puerta peatonal ni el cristal fijo. Será necesario colocar perfiles de metal en la parte superior e inferior para asegurar las duelas. El trabajo también abarca la colocación y la aplicación de barniz.Detalles del área: Área total: ______m x ______m, Área puerta peatonal: ______ m x ______m"
    body_text_bastidores_metalicos_duelasplastico="Este proyecto implica el suministro e instalación de tres puertas corredizas para cubrir un área total de ________ x ___________ metros en la puerta de acceso vehicular, así como una puerta peatonal abatible. Estas puertas se fabricarán utilizando bastidores de perfil metálico y se revestirán con duelas de plástico con un acabado que imita la apariencia de la madera. La instalación incluye guías y carretillas, tejuelos, y se llevará a cabo un trabajo de pintura antioxidante. Dimensiones del área de acceso vehicular: ____ x _________ metros  Dimensiones del área de la puerta peatonal: n x m"
    body_text_BARRERA_VEHICULAR_CLASSIC_SEG_24VDC   = "Te ofrecemos el suministro e instalación de la barrera automática SEG 24VDC CON LUZ LED, ideal para el control de acceso vehicular, conjugando eficiencia y seguridad en tus instalaciones. Este sistema, alimentado a 127 VAC/230 VAC y equipado con un motor de 24VDC, se distingue por su eficaz potencia de 1/2HP y un moderado consumo de 300W, asegurando una operación continua. Con una pluma de 3.3 metros y tiempos de apertura y cierre de apenas 4 segundos, garantiza un tránsito fluido y seguro. Se complementa con una lectora de tarjetas con capacidad para 500 usuarios y fotoceldas que previenen colisiones. El kit incluye 100 tarjetas y la configuración inicial del sistema, así como fotoceldas de seguridad, lectora de tarjetas y tarjeta controladora para dos accesos."
    body_text_mantenimiento_camaras = "El servicio de mantenimiento para las _____ cámaras previamente instaladas comprende la revisión de los conectores de alimentación y señales, el ordenamiento y ajuste del cableado, y la verificación de la alimentación de las cámaras para mejorar la calidad de la imagen de los dispositivos. Para lograr esto, es esencial la incorporación de baluns. Además, proporcionamos los conectores de alimentación y los baluns necesarios para la conexión de las cámaras. Estos dispositivos transforman señales no balanceadas en señales balanceadas, lo que asegura una transmisión óptima a través de cables UTP. También ofrecemos el suministro e instalación de una fuente de alimentación que garantiza que todas las cámaras reciban el nivel de voltaje adecuado, evitando así posibles distorsiones en la imagen debido a caídas de voltaje"
    body_text_puerta_peatonal1_ventana_herreria = "Seguridad sin privar tu casa de luz del día, ofrecemos suministro, diseño e instalación de puerta de acceso peatonal para la entrada de hogar. Con dimensiones  de 0.78__x___1.84 metros, fabricada a partir de un bastidor externo de perfil angular, un marco superior que alberga 6 cristales, y una guarda con perfiles cuadrados macizos de 1/4 de pulgada. El trabajo que realizamos incluye un acabado estético y trabajo de pintura. En la parte inferior de la puerta, instalamos una sólida calamina de calibre 18 para mayor durabilidad y resistencia. Además, nuestro precio también cubre una chapa de marca Phillips, una jaladera y un pasador de seguridad. Como un beneficio especial, en esta ocasión, Trix te obsequia un render con el diseño de tu puerta. Normalmente, este servicio de diseño tiene un costo adicional." 
    body_text_Mtto_Piston_EM  ="Servicio de mantenimiento para un par de pistones electromecánicos. El servicio incluye lo siguiente: limpieza y lubricación del ensamble de Engranaje Anti-retorno Helicoidal, además de la limpieza y lubricación del Tornillo de Avance. Adicionalmente, se reprogramará el recorrido para asegurar el correcto cierre de la puerta. "
    body_text_afpam1500_cremallera = "Suministro e instalación de operador de la marca Afpam, modelo 1500 a 110Vac, con capacidad para 1500 kg. Diseñado para automatizar puertas corredizas con mecanismo de piñón y cremallera. Cuenta con control de aceleración para la apertura y cierre."
    body_text_s800h_indv = "Suministro e instalación de un motor hidráulico subterráneo, marca FAAC, modelo S800H ENC 24V de 110 Vac de alimentación, apto para puertas de hasta 800 kg y 4 metros de largo. El paquete incluye un operador S800H, una caja de cimentación, y una caja electrónica para el control de operadores, sensores y accesorios. Además, incluye un botón para apertura desde el interior y dos controles de radiofrecuencia"
    body_text_s800h_par_kit = "Kit de suministro e instalación de motores hidráulicos subterráneos, marca FAAC, modelo S800H ENC 24V de 110 Vac de alimentación, diseñado para puertas de hasta 800 kg y 4 metros de largo. El kit incluye dos operadores S800H, cajas de cimentación y cajas electrónicas para el control de los operadores, sensores y accesorios. Además, el kit contiene un botón para apertura desde el interior y dos controles de radiofrecuencia."
    body_text_770N = "Kit de suministro e instalación de motores hidráulicos subterráneos FAAC, modelo 770N 24V, diseñado para puertas de acceso vehicular. Este kit incluye dos operadores 770N, funcionando a un voltaje de motor de 24V con una tensión de alimentación de red de 110 Vac60 Hz.  Los operadores son aptos para puertas de hasta 500 kg de peso . El kit también incluye cajas de cimentación, cajas electrónicas para el control de los operadores, sensores y accesorios, así como un botón para apertura desde el interior y dos controles de radiofrecuencia"
    body_text_plegadiza_a_abatible = "Trabajo de herreria para convertir cambio de apertura de puerta, actualmente cuentan con una puerta plegadiza de cuatro bastidores. Es necesario formar 2 hojas de ___ x____ apartir de los bastidores actuales. Incluye cambio de tejuelos y rodamientos, trabajo de soldadura y pintura."
    body_text_receptor_metalico_liftmaster ="Suministro e instalación de receptor de radiofrecuencia marca LiftMaster, modelo RL8002E, para puertas automáticas. Este receptor es compatible con los transmisores LiftMaster con código Billion y Rolling Code, y puede trabajar en dos frecuencias distintas, 433 MHz o 418 MHz, mediante el cambio del módulo de radiofrecuencia. Su voltaje de operación tiene un rango de 9 a 24Vac/dc y de 24 a 30Vac/dc. "
    body_text_receptor_859lm = "Suministro e instalación de receptor de radiofrecuencia marca LiftMaster, modelo 850LM, ideal para puertas automáticas. Este receptor avanzado utiliza una radio multifrecuencia de banda estrecha en 310 MHz, 315 MHz y 390 MHz, que prácticamente elimina las interferencias de radio. Está equipado con la tecnología Security+ 2.0™ y cuenta con DIP encriptado para mayor seguridad. Su voltaje de operación es de 12-24 V CC o 12-24 V CA. Además, tiene capacidad para manejar hasta 50 controles remotos en el Canal 1, y 20 controles remotos tanto en el Canal 2 como en el Canal 3. Diseñado tanto para comunidades cerradas como para edificios comerciales, el receptor universal y los controles remotos que lo acompañan ofrecen un alcance y seguridad superiores."
    body_text_mtto_cabezal = "Servicio de mantenimiento y reparación de puerta con cabezal Merik. Incluye limpieza del mecanismo con el desengrasante dieléctrico WD-40, ideal para estos sistemas, cambio y ajuste de banda dentada, programación de parámetros de apertura y cierre, acomodo de cableado y conexiones. Además, se revisará y ajustará el sensor de presencia y los periféricos asociados para asegurar un funcionamiento óptimo y seguro."
    body_text_carro_collarin_511_711 = "Suministro e instalación de collarín para riel de puerta automática, componente diseñado para desplazarse a lo largo del riel, impulsar la puerta y hacer parte del sistema de liberación manual de la puerta. "
    body_text_merik_ks3000_cabezal = "Suministro e instalacion de cabezal para puerta automática marca Merik modelo ks3000 con motor a 24Vdc sin escobillas, incluye perfil de soporte de motor, fuente de alimenteacion con electronica de control "
    body_text_merik_gilgenSLA_cabezal = "Suministro e instalacion de cabezal para puerta automática marca Giglen modelo SLA con motor a 24Vdc sin escobillas, puede manejar dos hojas de hasta 120 kg cada una , incluye perfil de soporte de motor, fuente de alimenteacion con electronica de control "
    body_text_merik_doormakaba_cabezal = ""
    body_text_seg_bldc_cabezal = "Suministro e instalaacion "
    body_text_segrt800=  'Suministro e instalación de abre puertas residencial a base de cadena marca Seg modelo RT800 de 1/2 h.p. a 24 vcd. Alimentación a 115 vca. , monofásico con motor reductor de engranes y transmisión de cadena link-belt con corredera en el sentido de apertura de la puerta y brazos articulados de acople a la misma, además cuenta con una luz de encendido de 4 1/2  min. De apagado automático, módulo wifi, dos controles remotos para automóvil y un botón para abrir desde el interior .' 
    body_text_seg_peatonal = 'Suministro e instalación de operador SEG para puertas peatonales. El kit incluye electrónica de control, motor a 24V, herrajes y dos botones touchless. Este operador es adecuado para puertas abatibles al interior o exterior con uso continuo.'
    body_text_trabajo_pintura ="Se realizará el suministro y aplicación de tres capas de pintura ( 1 primer , 2 de color), en dos puertas plegadizas de herrería, utilizadas como acceso vehicular. Este proceso se llevará a cabo tanto en el interior como en el exterior de las puertas. La pintura a utilizarse será de color negro satinado."
    body_text_riel_1400= "Suministro e instalación de _____ rieles modelo 1400 y soportes adicionales angulares, incluyendo el reemplazo de ___ carretillas y cuatro ruedas de 4inch de diámetro con rodamientos. "
    body_text_riel_1500= "Suministro e instalación de _____ rieles modelo 1500 y soportes adicionales angulares, incluyendo el reemplazo de ___ carretillas y cuatro ruedas de 4inch de diámetro con rodamientos."
    body_text_carretillas = ""
    body_text_bisagra_escuadra = "Suministro e instalación de cuatro bisagras en forma de escuadra (dos superiores, dos inferiores) para puertas de acceso vehicular, elaboradas a partir de herrería."
    body_text_faac_412 = "Suministro e instalación de un sistema de apertura de puertas, equipado con brazos electromecánicos de la marca FAAC , modelo 412  a 110 V CA. El sistema incluye los siguientes componentes: un par de brazos electromecánicos, un centro de control (cerebro), un receptor de control remoto, dos controles remotos digitales para automóviles, un botón de apertura, una llave de desbloqueo, herrajes y todos los elementos necesarios para su correcto funcionamiento."
    body_text_stanley_continuio_duraglide ="Suministro e instalación de cabezal para automatizar puerta de cristal, instalada en un vano total de 4.04x2.56m, marca Stanley modelo Dura Glide, de alimentación a 110V AC, cuenta con un motor brushless a 24V DC, sistema de transmision por banda dendata y opción a batería de respaldo, sensor de presencia, freno manual de tenasa, botón de acción de pared, colocación de cristales de 9mm de espesor. Se instalarán 2 cristales en cada extremo sin movimiento, al centro se colocarán un par de hojas que abrirán en sentidos opuestos actuadas por el cabezal y su sistema de bandas. Este modelo se selecciona según el uso continuo que tendrá, siendo un acceso con alta frecuencia de operaciones por hora.Al elegir el Dura Glide, usted estará invirtiendo en una solución confiable y duradera que mejorará la eficiencia y la imagen. El cristal estara colocado en un marco de aluminio con apariencia metalica, intermente se reforzará con fierro del se sujetaran las carretillas."
    body_text_fd20 = "Suministro e instalacion de operador Gilgen modelo FD20, para puertas peatonales, consta de montaje, electronica, uso continuo, bidireccional"
    body_text_SLF100_pivus_seg =" Abre puertas seg - pivus slf100"
    body_text_seg_slider_cabezal =" Suministro e instalación de cabezal para automatizar puerta de cristal ( cristal proporcionado por el cliente), instalada en un vano total de 4.04x2.56m, marca Seg modelo Slider BLDC, de alimentación a 110V AC, cuenta con un motor brushless a 24V DC, sistema de transmision por banda dendata y opción a batería de respaldo, sensor de presencia,  botón de acción de pared. Se instalarán  sobre  2 cristales ( proporcionado por el cliente) en cada extremo sin movimiento, al centro se colocarán un par de hojas que abrirán en sentidos opuestos actuadas por el cabezal y su sistema de bandas."
    body_PPA_cabezal_RAC = "Suministro e instalación de operador automatico MARCA PPA modelo RAC para puerta corrediza de cristal, de alimentación electrica @ 110VAC 60HZ. El operador cuenta con una capacidad máxima  de  la hoja carga de 80kg, recomendado en para puertas con una frecuencia de uso menor a 60 ciclos por hora. Este precio cuenta con el suministro y colocacion de hoja de cristal de 9mm , el cual estará automatizado por el motor. Cuenta con boton touchless, electrónica de control, herrajes y todo lo necesario para su correcto funcionamiento."
    body_bateria_711 = "SUministro e instalacion de batería de respaldo para operador Merik 711, esta bateria entra en operación en corte de suministro eléctrico"
    body_text_nota_recibo="Concepto: Servicio de mantenimiento para puerta automática y trabajo de soldadura, Direccion : Cda. Trabajadores Sociales 32, Aculco, Ciudad de México, Precio:  por la cantidad de $1,700.00 MXN, (mil setecientos pesos mexicanos)"
    body_resorte_izq ="Suministro e instalación de resorte izquierdo de torsión para puerta con mecanismo ascendente."
    body_text_tarjeta_brain_15 = " Suministro e instalación de tarjeta de control electrónica Genius, para operadores a 24VDC, incluye instalación y puesta en marcha del equipo."
    body_text_engrane_helicoidal_liftmaster = "Suministro e instalación de engrane helicoidal para ensamble de moto reductor Liftmaster, incluye servicio de mantenimiento, ajuste de cadena, cable de acero y puesta en marcha del equipo "
    body_text_fort_1000_comunello= "Suministro e instalación de motor a 24VDC para puerta corrediza con mecanismo piñón cremallera. Cuenta con 2 controles, 4 metros de cremallera, torreta indicadora con luz, baterias de respaldo, puesta en marcha de equipo."
    body_text_puerta_peaton_motor = "Suministro e instalación de tres puertas abatibles para cubrir un vano total de 0.9 x 2.2 m. Cada puerta estará fabricada con cristal de 6 mm y contará con un marco de aluminio. El servicio incluye la instalación del marco o jamba, bisagra de suelo y jaladera. Para automatizar, se instalara un operador automático marca SEG modelo SLF de 110VAC de alimentación y para puertas con peso máximo de 150 kg , en cada uno de los tres accesos. Incluye electronica de control, motor a 24VDC y 3A, botones touchless para operar puerta desde el interior y exterior. Se requiere de una fuerte estructura para el montaje de operadores y cristales.  "
# Crear formato de párrafo para la introducción
    intro_format = document.styles.add_style('Intro', WD_STYLE_TYPE.PARAGRAPH)
    intro_format.font.name = 'Arial'
    intro_format.font.size = Pt(12)
    intro_format.font.bold = True
    
# Crear formato de párrafo para nombre
    nombre_format = document.styles['Normal']
    nombre_format.font.name = 'Arial'
    nombre_format.font.size = Pt(10)
    nombre_format.font.bold = False

    dir_date_format = document.styles['Normal']
    dir_date_format.font.name = 'Arial'
    dir_date_format.font.size = Pt(10)
    dir_date_format.font.bold = False

    body_format = document.styles['Normal']
    body_format.font.name = 'Arial'
    body_format.font.size = Pt(10)
    body_format.font.bold = False

    add_fecha = document.add_paragraph(fecha_cotizacion + ',' + Ciudad_TRIX)
    add_fecha.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    add_direccion = document.add_paragraph(Direccion_cliente , style=nombre_format)

    add_direccion.paragraph_format.space_after = Pt(0)
    add_nombre = document.add_paragraph(Nombre_cliente, style=dir_date_format)

    add_intro = document.add_paragraph(style='Intro')
    add_intro.add_run(intro_text)
    

    Precio_text= 'Precio: $ + iva'


        # Extraer todos los productos seleccionados usando getlist
    productos_form = form_data.getlist('productos')
    logging.info(f"Productos seleccionados: {productos_form}")
    Cantidad_productos= len(productos_form)


#-------- bucle for para crear parrafos en función de variable número de productos

 # Agregar los productos al documento
    for i in range(Cantidad_productos):
        Producto = productos_form[i]
        Producto_text = f"Opción {i+1}: {Producto} "
        add_producto = document.add_paragraph(Producto_text)
        image_path = ''

        if Producto == 'Power230':
            body_text = body_text_power230
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'power230.jpg')
            
        elif Producto == 'gbat300':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'gbat300.jpg')
        
        elif Producto == 'nota_recibo':  
            body_text = body_text_nota_recibo

        elif Producto == 'SEGDuoCombat':
            body_text = body_text_SEGDuoCombat
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'segcombatduo.jpg')

        elif Producto == 'Pivus30024v':
            body_text = body_text_pivus30024v
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'pivus300.jpg')

        elif Producto == 'Faac_400_CBACL_kit':
            body_text = body_text_Faac_400_CBACL_kit
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'FAAC_400_CBAC.jpg')

        elif Producto == 'Faac_350h_kit':
            body_text = body_text_Faac_350h
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'FAAC_400_CBAC.jpg') 
        
        elif Producto == 's800h_indv':
            body_text = body_text_s800h_indv
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'S800H.jpg')
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'faac_sub_install.png')

        elif Producto == 's800h_par_kit':
            body_text = body_text_s800h_par_kit
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'par_s800h.png')


        elif Producto == 'Gbat110v':
            body_text = body_text_gbat110
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'gbat300.jpg')
        
        
        elif Producto == 'faac_412_brazos110v_EM':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'faac_412_EM_110v.png') 
            body_text =  body_text_faac_412 

        elif Producto == 'Gbat24v':
            body_text =    body_text_gbat24v 
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'gbat24v.jpg')

        elif Producto == 'Motor711':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', '711.jpg')
            body_text=body_text_711

        elif Producto == 'Motor511':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', '511.jpg') 
            body_text=body_text_511

        elif Producto == 'Merik411':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'merik411.jpg') 
            body_text=body_text_merik411
        
        elif Producto == 'Merikmyq':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'myq_liftmaster.png') 
            body_text=body_text_merikMYQ
        

        elif Producto == 'SEG_cadena_rt800':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'segrt800.png')
            body_text=body_text_segrt800
  
        elif Producto == 'SEG_operador_peatonal':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'seg_peatonal.png')
            body_text= body_text_seg_peatonal 
        
        elif Producto == 'Peatonal_motorSEG':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'seg_peatonal.png')
            body_text= body_text_puerta_peaton_motor
        
        

        elif Producto == 'Bateria_711':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'bateria711resp.png')
            body_text=body_bateria_711

            
        elif Producto == 'Resorte_torsion_IZQ':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'resorte_torsion.png')
            body_text=body_resorte_izq

        elif Producto == 'bmt5011':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'bmt5011.png')
            body_text=body_text_5011

        elif Producto == 'BARRERA_VEHICULAR_CLASSIC_SEG_24VDC':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'barrera_Seg_3m.jpg') 
            body_text=  body_text_BARRERA_VEHICULAR_CLASSIC_SEG_24VDC  


        elif Producto == 'dormakaba_80kg_cierrapuerta':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'dormakaba_80kg_cierrapuerta.jpg')
            body_text= boy_text_doormakaba80kg 

        elif Producto == 'mtto_cabezal':
                    body_text = body_text_mtto_cabezal  

        elif Producto == 'gilgen_sla':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'gilgensla.png')
            body_text= body_text_merik_gilgenSLA_cabezal

        elif Producto == 'merik_ks3000':
            body_text = body_text_merik_ks3000_cabezal
        
        elif Producto == 'seg_slider_cabezal':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'sliderbldc_seg-removebg.png')
            body_text= body_text_seg_slider_cabezal

        elif Producto == 'Stanley_duraglide':
            body_text = body_text_stanley_continuio_duraglide
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'stanley_slidewood.png')
            
        elif Producto == 'gilgel_fd20_peatonal':
            body_text = body_text_fd20
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'gilgen_fd20.png')
  
        elif Producto == 'PPA_cabezal_RAC':
            body_text = body_PPA_cabezal_RAC
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'PPA_RAC_CABEZAL_S.png')
            
        elif Producto == 'Roller24V':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'roller24v.jpg')
            body_text = body_text_roller

        elif Producto == 'Trabajo_de_Herreria':
            body_text = body_text_trabajo_herreria  
        
        elif Producto == 'Trabajo_de_pintura':
            body_text = body_text_trabajo_pintura

        elif Producto == 'Riel_1400_carretillas':
                    body_text = body_text_riel_1400

        elif Producto == 'Riel_1500_carretillas':
                    body_text = body_text_riel_1500

        
        elif Producto == 'plegadiza_a_abatible':
            body_text = body_text_plegadiza_a_abatible
            

        elif Producto == 'Seg_2puertas_controlacceso':
            body_text = body_text_Seg_2puertas_controlacceso
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'seg_control_2puertas.jpg')

        elif Producto == 'Shocker':
            body_text = body_text_shoker 
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'kit_shocker.jpg')

        elif Producto == 'videoporterohikvision':
            body_text = body_text_videoporterohikvision
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'videoportero_hikvision_Terminal_Min_Moe.jpg')
        
        elif Producto == 'Dahua_2mp_domo_2.8mm':
            body_text = body_text_Dahua_2mp
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'Dahua_domo_2mp.jpg')

        elif Producto == 'DVR16chanDAHUA':
            body_text = body_text_DVR16chan
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'switch16chan.jpg')

        elif Producto == 'Dahua_ptz':
            body_text = body_text_Dahua_PZT
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'Dahua_ptz.jpg')
               
        elif Producto == 'Mantenimiento_hidraulico':
            body_text = body_text_mantenimiento_hidraulico    

        elif Producto == 'Mtto_Piston_EM':
                    body_text = body_text_Mtto_Piston_EM    

        
        elif Producto == 'Mantenimiento_cerca':
            body_text = body_text_mantenimiento_cerca    

        elif Producto == 'Mantenimiento_camaras':
            body_text = body_text_mantenimiento_camaras

        elif Producto == 'ChapaBFT110v':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'bft110v.jpg')
            body_text=body_text_bft110v

        elif Producto == 'Tensor_motorcadena':
            body_text=body_text_Tensor_motorcadena

        elif Producto == 'Chapa24v':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'chapa_24vseg.jpg')
            body_text=body_text_chapa24v

        elif Producto == 'fococelda_power230':
                    image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'fococelda_power230.jpg')
                    body_text=body_text_fococelda_power230 

        elif Producto == 'Bateria12vSEG':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'pila12vseg.jpg')
            body_text=body_text_pila12v

        elif Producto == 'DashCAM_hikvision':
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'hikvisiondashcam.jpg')
            body_text=body_text_dashcam

        elif Producto == 'Tarjeta_electronica_mec200':
            body_text=body_text_tarjeta_mec200
            
        elif Producto == 'Tarjeta_genius_brain_15_24V':
            body_text= body_text_tarjeta_brain_15

    
        elif Producto == 'Cremallera':
            body_text=body_text_cremallera
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'cremallera.jpg')

        elif Producto == 'afpam1500_cremallera':
            body_text=body_text_afpam1500_cremallera
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'afpam1500_cremallera.png')

        elif Producto == 'Fotoceldas':
            body_text = body_text_fotocelda

        elif Producto == 'CadenaCh':
            body_text = body_text_cadenaCh

        elif Producto == 'CadenaG':
            body_text = body_text_cadenaG
        
        elif Producto == 'Tejueloybibel':
            body_text = body_text_tejuelosybibel
            
        elif Producto == 'Carretillas':
            body_text = body_text_carretilla
    
        elif Producto == 'Engrane_helicoidal':
            body_text =  body_text_engrane_helicoidal_liftmaster 

        elif Producto == 'bisagra_escuadra':
            body_text = body_text_bisagra_escuadra
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'render_bisagra_escuadras.jpg')

        
        elif Producto == 'Duelas_madera':
            body_text = body_text_duelas_madera

        elif Producto == 'bastidores_metalicos_duelasplastico':
            body_text = body_text_bastidores_metalicos_duelasplastico

        elif Producto == 'puerta_peatonal1_ventana_herreria_Sencilla':
            body_text = body_text_puerta_peatonal1_ventana_herreria
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'render_peatonal_herreria_ventana.jpg')

        elif Producto == 'Mantenimiento':
            body_text = body_text_mantenimiento
        
        elif Producto == 'receptor__433':
            body_text = body_text_receptor433hz
        
        elif Producto == 'receptor_administrableSEG':
            body_text = body_text_receptoradministrableSEG
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'receptor_administrableSEG.jpg')
        

        elif Producto == 'receptor_metalico_liftmaster':
            body_text = body_text_receptor_metalico_liftmaster
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'receptor_metalico_liftmaster.jpeg')

       
        elif Producto == 'fort_1000_cremallera_24v':

            body_text= body_text_fort_1000_comunello
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'fot1000_comunello.png')


        elif Producto == 'control_llavero_433':
            body_text = body_text_contol_llavero_433
            image_path = os.path.join(current_app.root_path, 'static', 'img', 'productos', 'control_433_llavero.jpg')
 
        else:
            image_path = '' 
            ficha_tecnica = "Ficha técnica no disponible para este producto"
        
       
        add_body = document.add_paragraph(body_text, style=body_format)
        add_body.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if image_path:
             
            image_paragraph = document.add_paragraph()
            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            image = image_paragraph.add_run()
            image.add_picture(image_path, width=Inches(2.5))

        add_precio = document.add_paragraph(Precio_text)
        add_precio.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

     # Crear un párrafo para los términos
    terms_paragraph = document.add_paragraph()
    # Agregar el texto de términos al párrafo
    
    terms_paragraph.add_run('Términos\n').bold = True
    terms_paragraph.add_run('1. Se nos proporcionará una toma de corriente a 110 vca. Cerca del aparato.\n')
    terms_paragraph.add_run('2. Cada control remoto adicional de un botón tiene un costo de $ 900.00 más iva.\n')
    terms_paragraph.add_run('3. Forma de pago 60% de anticipo y 40% al término del trabajo.\n')
    terms_paragraph.add_run('4. Tiempo de entrega 8 días a partir de la fecha de autorización y liquidación del anticipo.\n')
    terms_paragraph.add_run('5. En este precio se incluye instalación y garantía de servicio por un año.\n')
    terms_paragraph.add_run('6. Cotización con vigencia de 10 días.\n')
    terms_paragraph.add_run('7. Trabajos que requieran pintura en general, tienen un costo adicional.\n')
    terms_paragraph.add_run('8. Garantía: Para hacer efectiva la garantía, será necesario presentar la nota, factura o comprobante de pago correspondiente a las cuentas de TRIX.\n')
  
    # Footer

    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    footer.center_footer = True
    footer_text = footer.add_paragraph()
    footer_text.add_run('Puertas Automáticas TRIX')
    

 # Crear un párrafo para la despedida y firma
    farewell_paragraph = document.add_paragraph()
    # Agregar la despedida y firma al párrafo
    farewell_paragraph.add_run('\nAtte.\nIng. José Luis Huidobro Echeverría')
    farewell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    unique_id = uuid.uuid4().hex
    # Guardar el archivo
    ruta_archivo = os.path.join(current_app.root_path, 'static', 'cotizaciones', f"{fecha_cotizacion}-{Direccion_cliente}-{unique_id}.docx")
    
    # ******* AÑADE ESTAS DOS LÍNEAS *******
    directorio = os.path.dirname(ruta_archivo)
    os.makedirs(directorio, exist_ok=True)
# **************************************
    
    archivo=document.save(ruta_archivo)
       # Enviar el correo con el archivo adjunto
    clasificacion = predict_location(form_data.get('lat'),form_data.get('lon'))

    
    correo2(
        direccion=Direccion_cliente,
        lat=form_data.get('lat'),
        lon=form_data.get('lon'),
        nombre=Nombre_cliente,
        clasificacion=clasificacion,
        fecha=fecha_cotizacion,
        concepto_list=productos_form,
        archivo=ruta_archivo  # Pasar la ruta del archivo adjunto
    )
    
